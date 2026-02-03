"""
Document Parser - Extract text from PDF, DOCX, and plain text files.

Supports:
- PDF extraction using PyMuPDF (fitz)
- DOCX extraction using python-docx
- Plain text (.txt) files

This module extracts the references section and splits into individual entries.

v2.8.1: Added table content filtering to reduce false reference extractions.
"""

import re
import os
from dataclasses import dataclass, field
from typing import List, Optional, Tuple
from pathlib import Path


@dataclass
class DocumentContent:
    """Extracted content from a document."""
    file_path: str
    file_type: str  # "pdf", "docx", "txt"
    full_text: str
    references_section: str
    reference_entries: List[str] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)  # title, author if extractable
    extraction_warnings: List[str] = field(default_factory=list)


class DocumentParser:
    """
    Extract text from PDF, DOCX, or plain text files.
    Focuses on locating and parsing the references section.
    """
    
    # Headers that indicate the start of references section
    REFERENCE_HEADERS = [
        r"(?:^|\n)\s*references?\s*(?:\n|$)",
        r"(?:^|\n)\s*bibliography\s*(?:\n|$)",
        r"(?:^|\n)\s*works?\s+cited\s*(?:\n|$)",
        r"(?:^|\n)\s*literature\s+cited\s*(?:\n|$)",
        r"(?:^|\n)\s*cited\s+literature\s*(?:\n|$)",
        r"(?:^|\n)\s*sources?\s*(?:\n|$)",
    ]
    
    # Headers that indicate end of references (appendix, etc.)
    END_HEADERS = [
        r"(?:^|\n)\s*appendix",
        r"(?:^|\n)\s*appendices",
        r"(?:^|\n)\s*supplementary",
        r"(?:^|\n)\s*acknowledgements?",
        r"(?:^|\n)\s*author\s+contributions?",
        r"(?:^|\n)\s*conflict\s+of\s+interest",
        r"(?:^|\n)\s*funding",
        r"(?:^|\n)\s*ethics\s+statement",
    ]
    
    # Patterns that indicate table content (not references)
    TABLE_INDICATORS = [
        r"^\s*[\d.,]+\s*$",  # Just numbers (e.g., "3.5", "100")
        r"^\s*[\d.,]+\s*%\s*$",  # Percentages (e.g., "85.3%")
        r"^\s*[<>≤≥=±]\s*[\d.,]+",  # Statistical values (e.g., "<0.001", "±2.5")
        r"^\s*p\s*[<>=]\s*[\d.,]+",  # P-values (e.g., "p<0.05")
        r"^\s*\d+\s*/\s*\d+\s*$",  # Fractions (e.g., "3/10")
        r"^\s*\d+\s*-\s*\d+\s*$",  # Ranges without context (e.g., "10-15")
        r"^\s*n\s*=\s*\d+",  # Sample sizes (e.g., "n=50")
        r"^\s*CI\s*:?\s*[\d.,-]+",  # Confidence intervals
        r"^\s*OR\s*:?\s*[\d.,]+",  # Odds ratios
        r"^\s*HR\s*:?\s*[\d.,]+",  # Hazard ratios
        r"^\s*RR\s*:?\s*[\d.,]+",  # Risk ratios
        r"^\s*MD\s*:?\s*[\d.,-]+",  # Mean differences
        r"^\s*SMD\s*:?\s*[\d.,-]+",  # Standardized mean differences
        r"^\s*NNT\s*:?\s*[\d.,]+",  # Number needed to treat
        r"^\s*\(\s*[\d.,-]+\s*,\s*[\d.,-]+\s*\)\s*$",  # CI ranges like "(1.2, 3.4)"
        r"^\s*\[\s*[\d.,-]+\s*,\s*[\d.,-]+\s*\]\s*$",  # CI ranges like "[1.2, 3.4]"
        r"^\s*Yes\s*$",  # Binary values
        r"^\s*No\s*$",
        r"^\s*N/A\s*$",
        r"^\s*NA\s*$",
        r"^\s*NR\s*$",  # Not reported
        r"^\s*[-–—]+\s*$",  # Just dashes
        r"^\s*[✓✗×•·]+\s*$",  # Checkmarks and bullets
    ]
    
    # Minimum characteristics for a valid reference
    MIN_REFERENCE_LENGTH = 40  # Characters
    MIN_WORD_COUNT = 6  # Words
    
    def __init__(self):
        self._fitz_available = None
        self._docx_available = None
    
    def _check_dependencies(self, file_type: str) -> Tuple[bool, str]:
        """Check if required dependencies are available."""
        if file_type == "pdf":
            if self._fitz_available is None:
                try:
                    import fitz
                    self._fitz_available = True
                except ImportError:
                    self._fitz_available = False
            if not self._fitz_available:
                return False, "PyMuPDF not installed. Run: pip install PyMuPDF"
            return True, ""
        
        elif file_type == "docx":
            if self._docx_available is None:
                try:
                    import docx
                    self._docx_available = True
                except ImportError:
                    self._docx_available = False
            if not self._docx_available:
                return False, "python-docx not installed. Run: pip install python-docx"
            return True, ""
        
        return True, ""
    
    def _is_table_content(self, text: str) -> bool:
        """
        Detect if text looks like table data rather than a reference.
        
        Returns True if the text appears to be table content.
        """
        text = text.strip()
        
        # Empty or very short text is likely table content
        if len(text) < 10:
            return True
        
        # Check against table indicator patterns
        for pattern in self.TABLE_INDICATORS:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        # High proportion of numbers suggests table content
        digit_count = sum(1 for c in text if c.isdigit())
        if len(text) > 0 and digit_count / len(text) > 0.5:
            return True
        
        # Very few alphabetic characters
        alpha_count = sum(1 for c in text if c.isalpha())
        if len(text) > 5 and alpha_count / len(text) < 0.3:
            return True
        
        return False
    
    def _is_valid_reference(self, text: str) -> bool:
        """
        Validate that text looks like a legitimate reference.
        
        Checks:
        - Minimum length
        - Contains author-like patterns
        - Contains year
        - Not just table data
        """
        text = text.strip()
        
        # Check minimum length
        if len(text) < self.MIN_REFERENCE_LENGTH:
            return False
        
        # Check minimum word count
        words = text.split()
        if len(words) < self.MIN_WORD_COUNT:
            return False
        
        # Filter out table content
        if self._is_table_content(text):
            return False
        
        # Must contain a year (1900-2099)
        if not re.search(r'\b(19|20)\d{2}\b', text):
            return False
        
        # Should contain author-like pattern: "Name," or "Name, I." or "et al"
        author_patterns = [
            r'[A-Z][a-z]+,\s*[A-Z]\.?',  # Smith, J.
            r'[A-Z][a-z]+\s+[A-Z]\.',     # Smith J.
            r'et\s+al\.?',                 # et al.
            r'[A-Z][a-z]+,\s+[A-Z][a-z]+', # Last, First
        ]
        has_author = any(re.search(p, text) for p in author_patterns)
        if not has_author:
            return False
        
        # Should NOT be primarily a table header or column name
        table_header_patterns = [
            r'^(Study|Author|Year|Design|N|Sample|Outcome|Result|Intervention|Control|Mean|SD)\s*$',
            r'^Table\s+\d+',
            r'^Figure\s+\d+',
        ]
        for pattern in table_header_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return False
        
        return True
    
    def _filter_table_entries(self, entries: List[str]) -> Tuple[List[str], List[str]]:
        """
        Filter out entries that look like table content.
        
        Returns:
            Tuple of (valid_entries, filtered_entries)
        """
        valid = []
        filtered = []
        
        for entry in entries:
            if self._is_valid_reference(entry):
                valid.append(entry)
            else:
                filtered.append(entry)
        
        return valid, filtered
    
    def parse(self, file_path: str) -> DocumentContent:
        """
        Parse a document and extract references.
        
        Args:
            file_path: Path to PDF, DOCX, or TXT file
            
        Returns:
            DocumentContent with extracted text and references
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            file_type = "pdf"
        elif suffix in [".docx", ".doc"]:
            file_type = "docx"
        elif suffix in [".txt", ".text", ".md"]:
            file_type = "txt"
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
        
        # Check dependencies
        available, error_msg = self._check_dependencies(file_type)
        if not available:
            raise ImportError(error_msg)
        
        # Extract text based on file type
        if file_type == "pdf":
            full_text, metadata = self._extract_pdf(file_path)
        elif file_type == "docx":
            full_text, metadata = self._extract_docx(file_path)
        else:
            full_text, metadata = self._extract_txt(file_path)
        
        # Find references section
        references_section, warnings = self._find_references_section(full_text)
        
        # Split into individual entries
        raw_entries = self._split_references(references_section)
        
        # Filter out table content (v2.8.1)
        reference_entries, filtered_entries = self._filter_table_entries(raw_entries)
        
        # Add warning if we filtered entries
        if filtered_entries:
            warnings.append(
                f"Filtered {len(filtered_entries)} non-reference entries "
                f"(likely table content). Kept {len(reference_entries)} valid references."
            )
        
        return DocumentContent(
            file_path=str(path.absolute()),
            file_type=file_type,
            full_text=full_text,
            references_section=references_section,
            reference_entries=reference_entries,
            metadata=metadata,
            extraction_warnings=warnings
        )
    
    def _extract_pdf(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from PDF using PyMuPDF."""
        import fitz  # PyMuPDF
        
        metadata = {}
        text_parts = []
        
        doc = fitz.open(file_path)
        
        # Extract metadata
        pdf_metadata = doc.metadata
        if pdf_metadata:
            if pdf_metadata.get("title"):
                metadata["title"] = pdf_metadata["title"]
            if pdf_metadata.get("author"):
                metadata["author"] = pdf_metadata["author"]
        
        # Extract text from all pages
        for page in doc:
            text = page.get_text()
            text_parts.append(text)
        
        doc.close()
        
        full_text = "\n".join(text_parts)
        return full_text, metadata
    
    def _extract_docx(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from DOCX using python-docx."""
        from docx import Document
        
        metadata = {}
        text_parts = []
        
        doc = Document(file_path)
        
        # Extract core properties as metadata
        try:
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
        except Exception:
            pass  # Metadata extraction is optional
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        full_text = "\n".join(text_parts)
        return full_text, metadata
    
    def _extract_txt(self, file_path: str) -> Tuple[str, dict]:
        """Extract text from plain text file."""
        metadata = {}
        
        # Try different encodings
        for encoding in ["utf-8", "utf-16", "latin-1", "cp1252"]:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    full_text = f.read()
                return full_text, metadata
            except UnicodeDecodeError:
                continue
        
        # Fallback with error handling
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            full_text = f.read()
        
        return full_text, metadata
    
    def _find_references_section(self, text: str) -> Tuple[str, List[str]]:
        """
        Find the references section in the document.
        
        Returns:
            Tuple of (references_text, warnings)
        """
        warnings = []
        text_lower = text.lower()
        
        # Find start of references section
        start_pos = -1
        for pattern in self.REFERENCE_HEADERS:
            match = re.search(pattern, text_lower, re.IGNORECASE | re.MULTILINE)
            if match:
                # Use the position in original text
                start_pos = match.end()
                break
        
        if start_pos == -1:
            warnings.append("Could not locate References section header")
            # Try to find numbered references pattern as fallback
            numbered_match = re.search(r"\n\s*\[1\]|\n\s*1\.\s+[A-Z]", text)
            if numbered_match:
                start_pos = numbered_match.start()
                warnings.append("Using numbered reference pattern as fallback")
            else:
                # Return empty if we can't find references
                return "", warnings
        
        # Find end of references section
        end_pos = len(text)
        remaining_text = text[start_pos:].lower()
        
        for pattern in self.END_HEADERS:
            match = re.search(pattern, remaining_text, re.IGNORECASE | re.MULTILINE)
            if match:
                end_pos = start_pos + match.start()
                break
        
        references_section = text[start_pos:end_pos].strip()
        
        if not references_section:
            warnings.append("References section appears to be empty")
        
        return references_section, warnings
    
    def _split_references(self, refs_text: str) -> List[str]:
        """
        Split references section into individual entries.
        
        Handles:
        - Numbered references: [1], 1., (1)
        - APA style (author-date)
        - Elsevier/academic format: Author, A., Author, B., Year. Title...
        - Hanging indent style
        """
        if not refs_text:
            return []
        
        entries = []
        
        # Try numbered patterns first: [1], [2], etc.
        numbered_bracket = re.split(r"\n\s*\[\d+\]\s*", refs_text)
        if len(numbered_bracket) > 5:  # Must have at least 5 to be real
            entries = [e.strip() for e in numbered_bracket if e.strip()]
            return entries
        
        # Try numbered with period: 1., 2., etc.
        numbered_period = re.split(r"\n\s*\d+\.\s+", refs_text)
        if len(numbered_period) > 5:
            entries = [e.strip() for e in numbered_period if e.strip()]
            return entries
        
        # Try numbered in parentheses: (1), (2), etc.
        # Need stricter matching - at least 10 entries to be considered valid
        numbered_paren = re.split(r"\n\s*\(\d+\)\s*", refs_text)
        if len(numbered_paren) > 10:
            entries = [e.strip() for e in numbered_paren if e.strip()]
            return entries
        
        # For Elsevier/academic format without numbering:
        # Use line-by-line approach - look for author pattern at start of lines
        lines = refs_text.split('\n')
        entries = []
        current_entry = []
        
        # Pattern for start of a new reference: AuthorName, Initial.
        new_ref_pattern = re.compile(r'^[A-Z][a-z]+(?:[-\'][A-Z][a-z]+)?,\s+[A-Z]\.')
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            # Check if this line starts a new reference
            if new_ref_pattern.match(line_stripped) and current_entry:
                # Save the previous entry
                entry_text = ' '.join(current_entry)
                if len(entry_text) > 30:  # Minimum reasonable reference length
                    entries.append(entry_text)
                current_entry = [line_stripped]
            else:
                current_entry.append(line_stripped)
        
        # Don't forget the last entry
        if current_entry:
            entry_text = ' '.join(current_entry)
            if len(entry_text) > 30:
                entries.append(entry_text)
        
        # If we found enough entries, return them
        if len(entries) > 3:
            return entries
        
        # Try splitting by blank lines (common in APA)
        blank_line_split = re.split(r"\n\s*\n", refs_text)
        if len(blank_line_split) > 3:
            entries = [e.strip() for e in blank_line_split if e.strip() and len(e.strip()) > 30]
            if len(entries) > 3:
                return entries
        
        # Fallback: Try APA pattern split
        apa_pattern = r"\n(?=[A-Z][a-z]+(?:[-'][A-Z][a-z]+)?,\s+[A-Z]\.)"
        apa_split = re.split(apa_pattern, refs_text)
        if len(apa_split) > 1:
            entries = [e.strip() for e in apa_split if e.strip()]
            return entries
        
        # Fallback: split by lines that look like they start a new reference
        # (start with author name pattern)
        lines = refs_text.split("\n")
        current_entry = []
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_entry:
                    entries.append(" ".join(current_entry))
                    current_entry = []
            elif re.match(r"^[A-Z][a-z]+(?:[-'][A-Z][a-z]+)?,", line):
                # Looks like start of new APA reference
                if current_entry:
                    entries.append(" ".join(current_entry))
                current_entry = [line]
            else:
                current_entry.append(line)
        
        if current_entry:
            entries.append(" ".join(current_entry))
        
        return entries
    
    def parse_batch(self, directory: str, pattern: str = "*.pdf") -> List[DocumentContent]:
        """
        Parse multiple documents from a directory.
        
        Args:
            directory: Directory path
            pattern: Glob pattern for files (default: "*.pdf")
            
        Returns:
            List of DocumentContent objects
        """
        from pathlib import Path
        
        dir_path = Path(directory)
        if not dir_path.is_dir():
            raise ValueError(f"Not a directory: {directory}")
        
        results = []
        for file_path in dir_path.glob(pattern):
            try:
                content = self.parse(str(file_path))
                results.append(content)
            except Exception as e:
                # Create a failed result
                results.append(DocumentContent(
                    file_path=str(file_path),
                    file_type=file_path.suffix.lstrip("."),
                    full_text="",
                    references_section="",
                    reference_entries=[],
                    extraction_warnings=[f"Failed to parse: {str(e)}"]
                ))
        
        return results
